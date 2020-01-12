# -*- coding: utf-8 -*-
"""
Created on Fri Dec 27 05:28:03 2019
@author: Shruti
"""

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
import docx
import random

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'
VALUE = WORD_NAMESPACE + 'val'
BULLETNUM = WORD_NAMESPACE + 'numId'

def chunks(l, n):
    for i in range(0, len(l), n):
        yield l[i:i + n]

def get_docx_text(path, options):
    #Take the path of a docx file as argument, return the text in unicode.
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    para = []
    for paragraph in tree.getiterator(PARA):
        alltext = [node.text
                 for node in paragraph.getiterator(TEXT)
                 if node.text]
        for bullet in paragraph.getiterator(BULLETNUM):
            if len(para) > 0:
                paragraphs.append(''.join(para))
                para = []
        para.append(''.join(alltext))
    paragraphs.append(''.join(para)) #for the last para to be appended
    return '\n\n'.join(paragraphs), [paragraphs[i:i+options+1] for i in range(0, len(paragraphs), options+1)]

def save_docx(df1, df2, df3, df4):
    try:
        for i in range(0,max(len(df1), len(df2), len(df3), len(df4)-1)):
            doc = docx.Document()
            heading = 'Practice Question Paper ' + str(i)
            doc.add_heading(heading, 0)
            doc.add_heading('Section: Name A', 1)
            try:
                for q in df1[i]:
                    doc.add_paragraph(q[0], style='List Number')
                    for j in q[1:]:
                        doc.add_paragraph(j, style='List Bullet')
                    doc.add_page_break()
            except Exception:
                pass
            try:
                for q in df2[i]:
                    doc.add_paragraph(q[0], style='List Number')
                    for j in q[1:]:
                        doc.add_paragraph(j, style='List Bullet')
                    doc.add_page_break()
            except Exception:
                pass
            doc.add_page_break()
            doc.add_heading('Section: Name B', 1)
            try:
                for q in df3[i]:
                    doc.add_paragraph(q[0], style='List Number')
                    for j in q[1:]:
                        doc.add_paragraph(j, style='List Bullet')
                    doc.add_page_break()
            except Exception:
                pass
            try:
                for q in df4[i]:
                    doc.add_paragraph(q[0], style='List Number')
                    for j in q[1:]:
                        doc.add_paragraph(j, style='List Bullet')
                    doc.add_page_break()
            except Exception:
                pass
            doc.save('path' + heading + '.docx')
    except Exception:
        return 0
    return 1

Q1, df1 = get_docx_text('bullet-wise formatted document full path and name',5) #5 is the number of options in this set of question paper per queston
Q2, df2 = get_docx_text('bullet-wise formatted document full path and name',5) #5 is the number of options in this set of question paper per queston
Q3, df3 = get_docx_text('bullet-wise formatted document full path and name',5) #5 is the number of options in this set of question paper per queston
Q4, df4 = get_docx_text('bullet-wise formatted document full path and name',5) #5 is the number of options in this set of question paper per queston

random.seed(17) #seed could be any integer
random.shuffle(df1)
random.shuffle(df2)
random.shuffle(df3)
random.shuffle(df4)

df1chunk = list(chunks(df1, 20)) #20 represents list-of-list's size
df2chunk = list(chunks(df2, 20)) #20 represents list-of-list's size
df3chunk = list(chunks(df3, 15)) #15 represents list-of-list's size
df4chunk = list(chunks(df4, 15)) #15 represents list-of-list's size

print(save_docx(df1chunk, df2chunk, df3chunk, df4chunk))